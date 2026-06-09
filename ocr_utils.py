# ╔══════════════════════════════════════════════════════════════╗
# ║ ocr_utils.py                                                 ║
# ║ v3.2.0 — feat #67: extract_anketa_fields() → (fields,msg,   ║
# ║          raw_text); raw_text нужен для ocr-preview           ║
# ║                                                              ║
# ║  • PDF: текст без OCR                                       ║
# ║  • DOCX/DOC: абзацы + ТАБЛИЦЫ (спец-парсер MTS)             ║
# ║  • JPG/PNG: OCR через easyocr (если есть)                   ║
# ║  • Табличные реквизиты читаем напрямую из docx.tables       ║
# ║  • Раздел 1.1–6 продолжаем разбирать по тексту              ║
# ╚══════════════════════════════════════════════════════════════╝

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Dict, Tuple

import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)

try:
    import easyocr
    _HAS_EASYOCR = True
except ImportError:
    _HAS_EASYOCR = False

# fix #13: ленивая инициализация — модель грузится только при первом OCR-запросе
_OCR_READER = None


def _get_ocr_reader():
    """Возвращает easyocr.Reader, загружая модель при первом вызове."""
    global _OCR_READER
    if not _HAS_EASYOCR:
        return None
    if _OCR_READER is None:
        logger.info("OCR: загрузка модели easyocr...")
        _OCR_READER = easyocr.Reader(['ru', 'en'], gpu=False)
        logger.info("OCR: модель загружена")
    return _OCR_READER


# ─── НОРМАЛИЗАЦИЯ ЗНАЧЕНИЙ ────────────────────────────────────────────────────

def _norm_phone(raw: str) -> str:
    """
    Нормализует телефон: оставляет только цифры и '+', убирает мусор.
    Если цифр меньше 7 — возвращает пустую строку (это не телефон).
    """
    cleaned = re.sub(r"[^\d+\-() ]", "", raw).strip()
    digits_only = re.sub(r"\D", "", cleaned)
    if len(digits_only) < 7:
        return ""
    return cleaned


def _norm_inn(raw: str) -> str:
    """
    Нормализует ИНН: оставляет только цифры.
    Допустимая длина ИНН — 10 (юр. лицо) или 12 (физ. лицо).
    """
    digits = re.sub(r"\D", "", raw)
    if len(digits) in (10, 12):
        return digits
    return ""


# ─── БАЗОВОЕ ИЗВЛЕЧЕНИЕ ТЕКСТА ───────────────────────────────────────────────

def _extract_text_pdf(path: str) -> str:
    try:
        parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                txt = page.extract_text() or ""
                parts.append(txt)
        return "\n".join(parts)
    except Exception as e:
        logger.warning("OCR: не удалось извлечь текст из PDF '%s': %s", path, e)
        return ""


def _is_text_pdf(path: str) -> bool:
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages[:2]:
                if (page.extract_text() or "").strip():
                    return True
        return False
    except Exception:
        return False


def _extract_text_image(path: str) -> str:
    reader = _get_ocr_reader()
    if reader is None:
        return ""
    # fix #66: защита от падения на повреждённых/неподдерживаемых изображениях
    try:
        result = reader.readtext(path, detail=0, paragraph=True)
        return "\n".join(result)
    except Exception as e:
        logger.warning("OCR: не удалось распознать изображение '%s': %s", path, e)
        return ""


def _normalize_text(text: str) -> str:
    return text.replace("\r", "\n").replace("\xa0", " ")


# ─── ПАРСИНГ ТАБЛИЦ DOCX (РЕКВИЗИТЫ) ────────────────────────────────────────

def _parse_docx_tables(path: str) -> Dict[str, str]:
    """
    Читает только таблицы DOCX и вытягивает:
      • applicant_full_name
      • applicant_inn          ← fix #65: новый паттерн
      • postal_address
      • project_name
      • contact_person
      • contact_phone          ← fix #65: точный regex + _norm_phone()
      • contact_email          ← fix #65: исправлен offset для «e-mail»
      • jobs_total, jobs_foreign
      • investment_total
      • object_composition
      • construction_start, operation_start
      • product_nomenclature
    """
    try:
        doc = Document(path)
        fields: Dict[str, str] = {}

        def clean_fio(val: str) -> str:
            v = val.strip()
            low = v.lower()
            if low.startswith("ф.и.о") or low.startswith("фио"):
                parts = v.split(":", 1)
                if len(parts) == 2:
                    v = parts[1].strip()
            return v

        # fix #65: точный поиск метки телефона — только «телефон» или «тел.:»
        _PHONE_RE = re.compile(
            r"(?:телефон|тел\.?\s*:)\s*([+\d][\d\s\-().+]{5,})",
            re.IGNORECASE,
        )

        # fix #65: regex для e-mail, нет проблемы с offset
        _EMAIL_RE = re.compile(
            r"e[\-\s]?mail\s*[:\s]\s*(\S+@\S+)",
            re.IGNORECASE,
        )

        # fix #65: поиск ИНН — строго 10 или 12 цифр после метки
        _INN_RE = re.compile(
            r"инн\s*[:\s]\s*(\d{10}|\d{12})",
            re.IGNORECASE,
        )

        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]

            i = 0
            while i < len(rows):
                row = rows[i]
                joined = " | ".join(row).lower()
                joined_raw = " | ".join(row)  # оригинал для regex

                if "заявитель (инвестор" in joined:
                    value_parts = []
                    for c in row:
                        if "заявитель (инвестор" not in c.lower():
                            if c.strip():
                                value_parts.append(c.strip())
                    j = i + 1
                    while j < len(rows) and "заявитель (инвестор" not in " | ".join(rows[j]).lower() \
                            and "почтовый и юридический адрес" not in " | ".join(rows[j]).lower():
                        for c in rows[j]:
                            ct = c.strip()
                            if ct:
                                value_parts.append(ct)
                        j += 1
                    if value_parts and "applicant_full_name" not in fields:
                        fields["applicant_full_name"] = " ".join(value_parts)
                    i = j
                    continue

                # fix #65: извлечение ИНН из таблицы
                if "инн" in joined:
                    m = _INN_RE.search(joined_raw)
                    if m and "applicant_inn" not in fields:
                        fields["applicant_inn"] = m.group(1).strip()
                    i += 1
                    continue

                if "почтовый и юридический адрес" in joined:
                    value_parts = []
                    for c in row:
                        if "почтовый и юридический адрес" not in c.lower():
                            if c.strip():
                                value_parts.append(c.strip())
                    if value_parts:
                        fields["postal_address"] = " ".join(value_parts)
                    i += 1
                    continue

                if "название проекта (краткое описание" in joined or "название проекта" in joined:
                    value_parts = []
                    for c in row:
                        if "название проекта" not in c.lower():
                            if c.strip():
                                value_parts.append(c.strip())
                    if value_parts:
                        fields["project_name"] = " ".join(value_parts)
                    i += 1
                    continue

                if "уполномоченное лицо по ведению проекта" in joined:
                    j = i + 1
                    while j < len(rows):
                        line = " | ".join(rows[j])
                        low_line = line.lower()
                        if "планируемое количество постоянных" in low_line \
                                or "планируемый объем инвестиций" in low_line:
                            break
                        if "ф.и.о" in low_line or "фио" in low_line:
                            val = clean_fio(line)
                            if val:
                                fields["contact_person"] = val

                        # fix #65: точный regex вместо жадного find("тел")
                        if "contact_phone" not in fields:
                            m_phone = _PHONE_RE.search(line)
                            if m_phone:
                                normed = _norm_phone(m_phone.group(1))
                                if normed:
                                    fields["contact_phone"] = normed

                        # fix #65: regex для e-mail, нет ошибки с offset
                        if "contact_email" not in fields:
                            m_email = _EMAIL_RE.search(line)
                            if m_email:
                                fields["contact_email"] = m_email.group(1).strip()

                        j += 1
                    i = j
                    continue

                if "планируемое количество постоянных рабочих мест" in joined:
                    line = " ".join(row)
                    digits = [
                        d for d in "".join(ch if ch.isdigit() else " " for ch in line).split()
                        if d.isdigit()
                    ]
                    if digits:
                        fields["jobs_total"] = digits[0]
                    if len(digits) > 1:
                        fields["jobs_foreign"] = digits[1]
                    i += 1
                    continue

                if "планируемый объем инвестиций" in joined or "планируемый объём инвестиций" in joined:
                    inv_lines = []
                    for c in row:
                        if "планируемый объем инвестиций" not in c.lower() \
                                and "планируемый объём инвестиций" not in c.lower():
                            if c.strip():
                                inv_lines.append(c.strip())
                    j = i + 1
                    while j < len(rows):
                        l = " | ".join(rows[j]).strip()
                        low_l = l.lower()
                        if not l:
                            j += 1
                            continue
                        if "описание строительства" in low_l or "планируемый срок начала строительства" in low_l:
                            break
                        inv_lines.append(l)
                        j += 1
                    for l in inv_lines:
                        if any(ch.isdigit() for ch in l):
                            fields["investment_total"] = l
                            break
                    i = j
                    continue

                if "описание строительства" in joined:
                    value_parts = []
                    for c in row:
                        if "описание строительства" not in c.lower():
                            if c.strip():
                                value_parts.append(c.strip())
                    j = i + 1
                    while j < len(rows):
                        l = " | ".join(rows[j]).strip()
                        if not l:
                            j += 1
                            continue
                        low_l = l.lower()
                        if "планируемый срок начала строительства" in low_l \
                                or "планируемый срок ввода предприятия" in low_l:
                            break
                        value_parts.append(l)
                        j += 1
                    if value_parts:
                        fields["object_composition"] = " ".join(value_parts)
                    i = j
                    continue

                if "планируемый срок начала строительства" in joined:
                    for c in row:
                        if "планируемый срок начала строительства" not in c.lower() and c.strip():
                            fields["construction_start"] = c.strip()
                    i += 1
                    continue

                if "планируемый срок ввода предприятия в эксплуатацию" in joined:
                    val = " ".join([c for c in row if "планируемый срок ввода предприятия" not in c.lower()]).strip()
                    if val:
                        fields["operation_start"] = val
                    i += 1
                    continue

                if "номенклатура планируемой к выпуску продукции" in joined:
                    val = " ".join([c for c in row if "номенклатура планируемой к выпуску продукции" not in c.lower()]).strip()
                    if val:
                        fields["product_nomenclature"] = val
                    i += 1
                    continue

                i += 1

        fields = {k: v for k, v in fields.items() if v}
        return fields
    except Exception as e:
        logger.warning("OCR: не удалось прочитать таблицы DOCX '%s': %s", path, e)
        return {}


# ─── ПАРСИНГ ТЕКСТА (РАЗДЕЛЫ 1.1–6, PDF и т.п.) ─────────────────────────────

def _parse_anketa_text_blocks(text: str) -> Dict[str, str]:
    raw = _normalize_text(text)
    low = raw.lower()

    def slice_block(start_anchor: str, stop_anchors) -> str:
        s_idx = low.find(start_anchor.lower())
        if s_idx == -1:
            return ""
        start = s_idx + len(start_anchor)
        end = len(raw)
        if isinstance(stop_anchors, str):
            stops = [stop_anchors]
        else:
            stops = stop_anchors
        for stop in stops:
            i = low.find(stop.lower(), start)
            if i != -1 and i < end:
                end = i
        return raw[start:end].strip(" \t:;–-\n")

    def find_after(labels, max_len: int = 120) -> str:
        if isinstance(labels, str):
            labels_list = [labels]
        else:
            labels_list = labels
        for label in labels_list:
            idx = low.find(label.lower())
            if idx != -1:
                start = idx + len(label)
                tail = raw[start:start+max_len]
                line = tail.splitlines()[0]
                return line.strip(" \t:;–-|")
        return ""

    fields: Dict[str, str] = {}

    prod_desc = slice_block(
        "1.1. краткое описание производства и используемых технологий",
        ["\n1.2.", "\nii.", "\n2.", "\nII "]
    )
    if prod_desc:
        fields["production_description"] = prod_desc

    if "object_composition" not in fields:
        comp = find_after(["1.2. состав объекта:", "1.2. состав объекта"], max_len=400)
        if comp:
            fields["object_composition"] = comp

    eng_extra = slice_block(
        "при необходимости укажите дополнительные требования к инженерной инфраструктуре:",
        ["\n2.", "\n2. ", "\n3.", "\n3. "]
    )
    if eng_extra:
        fields["engineering_extra"] = eng_extra

    transport_extra = slice_block(
        "2.3. при необходимости укажите дополнительные требования к транспортной инфраструктуре",
        ["\n3.", "\n3. "]
    )
    if transport_extra:
        fields["transport_extra"] = transport_extra

    # fix #65: извлечение preferred_districts из текстовых блоков PDF/DOCX
    districts = find_after(
        [
            "предпочтительный район размещения:",
            "предпочтительный район размещения",
            "предпочтительные районы:",
            "предпочтительные районы",
            "2.1. предпочтительный район",
            "район размещения:",
        ],
        max_len=300,
    )
    if districts:
        fields["preferred_districts"] = districts

    # fix #65: ИНН из текста (PDF-анкеты)
    inn_match = re.search(r"инн\s*[:\s]\s*(\d{10}|\d{12})", low)
    if inn_match and "applicant_inn" not in fields:
        fields["applicant_inn"] = inn_match.group(1)

    add_info = slice_block(
        "6. дополнительная информация:",
        []
    )
    if add_info:
        fields["additional_info"] = add_info

    return {k: v for k, v in fields.items() if v}


# ─── ПУБЛИЧНАЯ ФУНКЦИЯ ───────────────────────────────────────────────────────

def extract_anketa_fields(path: str) -> Tuple[Dict[str, str], str, str]:
    """
    На вход: путь к файлу анкеты.
    На выход:
      fields   — dict под ALL_FIELDS
      msg      — человекочитаемое описание, как обрабатывали файл
      raw_text — сырой извлечённый текст (для ocr-preview, #67)
    """
    p: Path = Path(path)
    ext = (p.suffix or "").lower()

    # fix #66: проверка существования и размера файла до любой обработки
    if not p.exists():
        logger.error("OCR: файл не найден '%s'", path)
        return {}, f"Файл не найден: {path}", ""
    if p.stat().st_size == 0:
        logger.error("OCR: файл пустой (0 байт) '%s'", path)
        return {}, "Файл пустой (0 байт) — загрузите корректный файл анкеты.", ""

    logger.debug("OCR: path=%s ext=%s", path, ext)

    text = ""
    msg = ""
    fields: Dict[str, str] = {}

    if ext in (".docx", ".doc"):
        try:
            table_fields = _parse_docx_tables(path)
            fields.update(table_fields)
            doc = Document(path)
            parts = []
            for pgh in doc.paragraphs:
                t = pgh.text.strip()
                if t:
                    parts.append(t)
            text = "\n".join(parts)
            msg = "Файл анкеты обработан как DOC/DOCX (таблицы + текст, без OCR)."
        except Exception as e:
            return {}, f"Не удалось прочитать DOC/DOCX: {e}", ""

    elif ext == ".pdf":
        if _is_text_pdf(path):
            text = _extract_text_pdf(path)
            msg = "Файл анкеты обработан как текстовый PDF (без OCR)."
        else:
            if not _HAS_EASYOCR:
                return {}, "Анкета похожа на скан PDF, но OCR (easyocr) на сервере не установлен.", ""
            return {}, "Сканированный PDF пока не поддерживается (нужна доработка OCR по картинкам).", ""

    elif ext in (".jpg", ".jpeg", ".png"):
        if not _HAS_EASYOCR:
            return {}, "Для обработки сканов анкеты нужен easyocr, который сейчас не установлен.", ""
        text = _extract_text_image(path)
        msg = "Файл анкеты обработан как изображение (OCR)."

    else:
        # fix #66: logger.error вместо молчаливого провала в else-блоке
        try:
            if _is_text_pdf(path):
                text = _extract_text_pdf(path)
                msg = "Файл анкеты обработан как PDF без расширения (эвристика по содержимому)."
            else:
                fields.update(_parse_docx_tables(path))
                doc = Document(path)
                parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                text = "\n".join(parts)
                msg = "Файл анкеты обработан как DOCX/DOC без расширения (эвристика по содержимому)."
        except Exception as e:
            logger.error("OCR: не удалось определить тип файла '%s': %s", path, e)
            return {}, (
                "Неподдерживаемый формат файла. Загрузите PDF или DOCX. Ошибка: " + str(e)
            ), ""

    if text and text.strip():
        logger.debug("OCR TEXT SAMPLE: %s", repr(text[:500]))
        block_fields = _parse_anketa_text_blocks(text)
        for k, v in block_fields.items():
            fields.setdefault(k, v)

    if not fields:
        return {}, msg or "Не удалось распознать структуру анкеты.", text

    return fields, msg, text
