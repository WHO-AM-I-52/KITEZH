from pathlib import Path
from docx import Document


def debug_docx_tables(path: str, max_tables: int = 5):
    p = Path(path)
    print(f"=== FILE: {p} ===")
    if not p.is_file():
        print("!!! NOT FOUND")
        return

    doc = Document(str(p))
    print(f"tables count: {len(doc.tables)}")

    for ti, table in enumerate(doc.tables):
        if ti >= max_tables:
            break
        print(f"\n--- TABLE {ti} ---")
        for ri, row in enumerate(table.rows):
            cells = [repr(cell.text.strip()) for cell in row.cells]
            print(f"row {ri}: {cells}")


if __name__ == "__main__":
    files = [
        r"C:\Users\aasoldatov\Downloads\Анкета_подбор площадки2222.docx",
        r"C:\Users\aasoldatov\Downloads\Анкета_подбор площадки МТС2.docx",
        r"C:\Users\aasoldatov\Downloads\АнкетаЦДС-НН.docx",
        r"C:\Users\aasoldatov\Downloads\Анкета заполненная_ООО Технолайн.docx",
        r"C:\Users\aasoldatov\Downloads\Анкета Калининград 39 плав средства.docx",
    ]

    for f in files:
        debug_docx_tables(f, max_tables=3)
        print("\n" + "=" * 80 + "\n")